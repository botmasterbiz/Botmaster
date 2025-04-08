import os
import argparse
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from crewai import Agent, Task, Crew, Process
from langchain_openai import ChatOpenAI

def create_toolbox_talk_form(transcript, output_file=None):
    """
    Create a toolbox talk form from a meeting transcript.
    
    Args:
        transcript (str): The meeting transcript text
        output_file (str, optional): Path to save the output Word document. 
                                    If None, a default name will be used.
    
    Returns:
        str: Path to the generated Word document
    """
    # Set up the LLM
    llm = ChatOpenAI(model="gpt-4o", temperature=0.2)
    
    # Create agents
    transcription_analyzer = Agent(
        role="Transcription Analyzer",
        goal="Extract key safety information from meeting transcripts",
        backstory="You are an expert at analyzing safety meetings and extracting critical information.",
        verbose=True,
        llm=llm
    )
    
    form_creator = Agent(
        role="Toolbox Talk Form Creator",
        goal="Create well-structured toolbox talk forms from meeting content",
        backstory="You specialize in organizing safety information into clear, actionable toolbox talk documentation.",
        verbose=True,
        llm=llm
    )
    
    # Define the tasks
    analyze_transcript_task = Task(
        description=f"""
        Analyze the following meeting transcript and extract:
        1. Key points related to safety or processes
        2. Safety reminders
        3. Recent incidents or concerns mentioned
        4. Questions and feedback from the meeting
        5. Action items assigned during the meeting
        
        Transcript:
        {transcript}
        
        Your output should be structured as a Python dictionary with keys for each category.
        """,
        agent=transcription_analyzer,
        expected_output="A dictionary containing structured extraction of meeting content."
    )
    
    create_form_task = Task(
        description=f"""
        Create a toolbox talk form using the following extracted content:
        {{extracted_content}}
        
        The form should include:
        1. Date (use current date)
        2. Space for facilitator name
        3. Space for time and location
        4. Space for department/team
        5. Topic of discussion derived from the key points
        6. The key points formatted as bullet points
        7. Safety reminders as bullet points
        8. Recent incidents/concerns as bullet points
        9. Questions & feedback as bullet points
        10. Action items as bullet points
        
        Format your response as a Python dictionary with keys for each section of the form.
        """,
        agent=form_creator,
        expected_output="A dictionary containing all content for a toolbox talk form."
    )
    
    # Create the crew
    crew = Crew(
        agents=[transcription_analyzer, form_creator],
        tasks=[analyze_transcript_task, create_form_task],
        process=Process.sequential,
        verbose=True
    )
    
    # Run the crew
    result = crew.kickoff()
    
    # Extract the form data from the result
    # The result will be a string representation of a Python dictionary
    # We need to evaluate it to get the actual dictionary
    try:
        # Convert the string to a dictionary
        form_data = eval(result)
    except Exception as e:
        print(f"Error parsing form data: {e}")
        # If we can't parse the result, create a default form
        form_data = {
            'date': datetime.now().strftime("%B %d, %Y"),
            'facilitator': '',
            'time': '',
            'location': '',
            'department': '',
            'topic': 'Safety Discussion',
            'key_points': ['No key points extracted'],
            'safety_reminders': ['No safety reminders extracted'],
            'incidents_concerns': ['No incidents or concerns extracted'],
            'questions_feedback': ['No questions or feedback extracted'],
            'action_items': ['No action items extracted']
        }
    
    # Generate the output filename if not provided
    if output_file is None:
        current_date = datetime.now().strftime("%B_%d_%Y")
        output_file = f"Toolbox_Talk_{current_date}.docx"
    
    # Create the Word document
    create_word_document(form_data, output_file)
    
    return output_file

def create_word_document(form_data, output_file):
    """Create a Word document based on the toolbox talk form data"""
    doc = Document()
    
    # Set up margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add title
    title = doc.add_heading("Toolbox Talk Form", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header information
    header_table = doc.add_table(rows=3, cols=2)
    header_table.style = 'Table Grid'
    
    # Row 1: Date and Facilitator
    header_table.cell(0, 0).text = f"Date: {form_data.get('date', '_________________')}"
    header_table.cell(0, 1).text = f"Facilitator: {form_data.get('facilitator', '__________________')}"
    
    # Row 2: Time and Location
    header_table.cell(1, 0).text = f"Time: {form_data.get('time', '_________________')}"
    header_table.cell(1, 1).text = f"Location: {form_data.get('location', '___________________')}"
    
    # Row 3: Department/Team (spans two columns)
    department_cell = header_table.cell(2, 0)
    department_cell.merge(header_table.cell(2, 1))
    department_cell.text = f"Department/Team: {form_data.get('department', '___________________')}"
    
    # Add topic of discussion
    doc.add_paragraph()
    topic_para = doc.add_paragraph("Topic of Discussion:")
    topic_para.style = 'Heading 2'
    doc.add_paragraph(form_data.get('topic', '')).italic = True
    doc.add_paragraph()
    
    # Add content sections
    sections = [
        ("1. Key Points Discussed:", "key_points"),
        ("2. Safety Reminders:", "safety_reminders"),
        ("3. Recent Incidents/Concerns:", "incidents_concerns"),
        ("4. Questions & Feedback:", "questions_feedback"),
        ("5. Action Items (if any):", "action_items")
    ]
    
    for section_title, key in sections:
        # Add section heading
        heading = doc.add_paragraph(section_title)
        heading.style = 'Heading 3'
        
        # Add bullet points
        items = form_data.get(key, [])
        if isinstance(items, str):
            # If items is a string, split it by newlines
            items = [item.strip() for item in items.split('\n') if item.strip()]
        
        for item in items:
            if item and item.strip():  # Check if item is not empty
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(item.strip())
        
        # Add space after section
        doc.add_paragraph()
    
    # Add attendees section
    doc.add_heading("Attendees (Print Name & Sign):", 1)
    
    # Create a table for attendees with 17 rows
    attendee_table = doc.add_table(rows=17, cols=2)
    attendee_table.style = 'Table Grid'
    
    # Fill in attendee rows
    for i in range(17):
        attendee_table.cell(i, 0).text = f"Name {i+1}: ____________________"
        attendee_table.cell(i, 1).text = f"Signature: ___________________"
    
    # Add facilitator signature at bottom
    doc.add_paragraph()
    signature_para = doc.add_paragraph("Facilitator Signature: ___________________ Date: __________________")
    signature_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Save the document
    doc.save(output_file)
    print(f"Document saved as {output_file}")

def main():
    parser = argparse.ArgumentParser(description='Generate a Toolbox Talk Form from meeting transcript')
    parser.add_argument('--transcript', type=str, required=True, help='Path to the meeting transcript file')
    parser.add_argument('--output', type=str, default='toolbox_talk_form.docx', help='Output filename for the Word document')
    args = parser.parse_args()
    
    # Read the transcript file
    try:
        with open(args.transcript, 'r') as file:
            transcript = file.read()
    except FileNotFoundError:
        print(f"Error: File '{args.transcript}' not found.")
        return
    except Exception as e:
        print(f"Error reading transcript file: {e}")
        return
    
    # Create the toolbox talk form
    output_file = create_toolbox_talk_form(transcript, args.output)
    print(f"Toolbox talk form created: {output_file}")

if __name__ == "__main__":
    main() 